import streamlit as st
from openai import OpenAI
import json
import io

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Resume Analyzer | Interview Helper",
    page_icon="📄",
    layout="wide",
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

/* Dark gradient background */
.stApp {
    background: linear-gradient(135deg, #050f1a, #0a2540, #0d1f2d);
    color: #e2e8f0;
}

/* Hide Streamlit chrome */
#MainMenu, footer, header { visibility: hidden; }

/* Top nav bar */
.nav-bar {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 0.6rem 1.2rem;
    background: rgba(255,255,255,0.05);
    border-bottom: 1px solid rgba(255,255,255,0.08);
    border-radius: 12px;
    margin-bottom: 1.5rem;
    backdrop-filter: blur(10px);
}

.nav-title {
    font-size: 1.3rem;
    font-weight: 700;
    background: linear-gradient(90deg, #2dd4bf, #38bdf8);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

/* Card style */
.glass-card {
    background: rgba(255,255,255,0.05);
    border: 1px solid rgba(255,255,255,0.1);
    border-radius: 16px;
    padding: 1.4rem 1.6rem;
    backdrop-filter: blur(12px);
    margin-bottom: 1rem;
}

/* Score circle */
.score-ring-wrap {
    display: flex;
    flex-direction: column;
    align-items: center;
    padding: 0.5rem 0;
}

.score-circle {
    width: 90px;
    height: 90px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 1.6rem;
    font-weight: 700;
    margin-bottom: 0.4rem;
    border: 4px solid;
}

.score-label {
    font-size: 0.72rem;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: #94a3b8;
    text-align: center;
}

/* Meter bar */
.meter-wrap { margin: 0.4rem 0 0.8rem 0; }
.meter-label {
    display: flex;
    justify-content: space-between;
    font-size: 0.82rem;
    color: #cbd5e1;
    margin-bottom: 4px;
}
.meter-track {
    height: 8px;
    background: rgba(255,255,255,0.1);
    border-radius: 99px;
    overflow: hidden;
}
.meter-fill {
    height: 100%;
    border-radius: 99px;
    transition: width 0.6s ease;
}

/* Badge */
.badge {
    display: inline-block;
    padding: 3px 10px;
    border-radius: 99px;
    font-size: 0.72rem;
    font-weight: 600;
    margin: 2px 3px;
}

/* Overall score banner */
.overall-banner {
    background: linear-gradient(135deg, rgba(20,184,166,0.2), rgba(56,189,248,0.2));
    border: 1px solid rgba(20,184,166,0.4);
    border-radius: 20px;
    padding: 1.8rem 2rem;
    text-align: center;
    margin-bottom: 1.5rem;
}

.overall-score-num {
    font-size: 4rem;
    font-weight: 800;
    background: linear-gradient(90deg, #2dd4bf, #38bdf8);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    line-height: 1;
}

.verdict-chip {
    display: inline-block;
    padding: 6px 18px;
    border-radius: 99px;
    font-size: 0.9rem;
    font-weight: 700;
    margin-top: 0.6rem;
}

.section-title {
    font-size: 1.05rem;
    font-weight: 700;
    color: #e2e8f0;
    margin-bottom: 0.8rem;
    border-left: 3px solid #14b8a6;
    padding-left: 0.6rem;
}

.reasoning-box {
    background: rgba(0,0,0,0.25);
    border-left: 3px solid #38bdf8;
    border-radius: 0 8px 8px 0;
    padding: 0.8rem 1rem;
    font-size: 0.87rem;
    color: #94a3b8;
    line-height: 1.6;
    margin-top: 0.5rem;
}

.tag-match { background: rgba(34,197,94,0.2); color: #4ade80; border: 1px solid rgba(34,197,94,0.3); }
.tag-partial { background: rgba(251,191,36,0.2); color: #fbbf24; border: 1px solid rgba(251,191,36,0.3); }
.tag-miss { background: rgba(239,68,68,0.2); color: #f87171; border: 1px solid rgba(239,68,68,0.3); }

/* Input area overrides */
.stTextArea textarea, .stTextInput input {
    background: rgba(255,255,255,0.06) !important;
    border: 1px solid rgba(255,255,255,0.15) !important;
    border-radius: 10px !important;
    color: #e2e8f0 !important;
}

.stFileUploader {
    border: 2px dashed rgba(20,184,166,0.45) !important;
    border-radius: 12px !important;
    background: rgba(20,184,166,0.05) !important;
}

/* Buttons */
.stButton > button {
    background: linear-gradient(135deg, #0d9488, #0369a1) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    padding: 0.55rem 1.6rem !important;
    transition: opacity 0.2s ease !important;
}

.stButton > button:hover {
    opacity: 0.88 !important;
}

/* Batch leaderboard table */
.batch-table {
    width: 100%;
    border-collapse: collapse;
    margin-bottom: 1.5rem;
}
.batch-table th {
    text-align: left;
    font-size: 0.75rem;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.07em;
    color: #64748b;
    padding: 0.5rem 0.8rem;
    border-bottom: 1px solid rgba(255,255,255,0.08);
}
.batch-table td {
    padding: 0.6rem 0.8rem;
    font-size: 0.88rem;
    border-bottom: 1px solid rgba(255,255,255,0.05);
    vertical-align: middle;
}
.batch-table tr:hover td { background: rgba(255,255,255,0.03); }
.rank-medal { font-size: 1.2rem; }
</style>
""", unsafe_allow_html=True)

# ── Helper: extract text from uploaded file ───────────────────────────────────
def extract_text(uploaded_file) -> str:
    """Return plain text from a .txt, .pdf, or .docx uploaded file."""
    if uploaded_file is None:
        return ""
    filename = uploaded_file.name.lower()
    raw = uploaded_file.read()

    if filename.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")

    elif filename.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            st.error("Install `pypdf` to parse PDFs: `pip install pypdf`")
            return ""

    elif filename.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            lines = []
            # Extract paragraphs (body text, headings, bullet points)
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())
            # Also extract text from tables (common in resume templates)
            for table in doc.tables:
                for row in table.rows:
                    row_text = "  ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        lines.append(row_text)
            return "\n".join(lines)
        except ImportError:
            st.error("Install `python-docx` to parse DOCX files: `pip install python-docx`")
            return ""

    else:
        st.warning("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")
        return ""

# ── Helper: colour for a score 0-10 ──────────────────────────────────────────
def score_color(score: float):
    if score >= 7.5:
        return "#4ade80", "#14532d"   # green text, dark green bg
    elif score >= 5:
        return "#fbbf24", "#451a03"   # amber text, dark amber bg
    else:
        return "#f87171", "#450a0a"   # red text, dark red bg

def meter_gradient(score: float) -> str:
    if score >= 7.5:
        return "linear-gradient(90deg, #22c55e, #4ade80)"
    elif score >= 5:
        return "linear-gradient(90deg, #f59e0b, #fbbf24)"
    else:
        return "linear-gradient(90deg, #ef4444, #f87171)"

def verdict_style(overall: float):
    if overall >= 7.5:
        return "background:#14532d; color:#4ade80;"
    elif overall >= 5:
        return "background:#451a03; color:#fbbf24;"
    else:
        return "background:#450a0a; color:#f87171;"

def verdict_label(overall: float) -> str:
    if overall >= 8:
        return "🌟 Excellent Fit"
    elif overall >= 7:
        return "✅ Strong Fit"
    elif overall >= 5:
        return "🟡 Moderate Fit"
    else:
        return "❌ Poor Fit"

# ── Prompt builder ────────────────────────────────────────────────────────────
ANALYSIS_SCHEMA = """{
  "overall_score": <float 0-10>,
  "overall_reasoning": "<string>",
  "dimensions": [
    {
      "name": "<dimension name>",
      "score": <float 0-10>,
      "matched": ["<keyword/skill>"],
      "partial": ["<keyword/skill>"],
      "missing": ["<keyword/skill>"],
      "reasoning": "<detailed reasoning string>"
    }
  ],
  "strengths": ["<string>"],
  "gaps": ["<string>"],
  "recommendation": "<Hire | Maybe | Pass>"
}"""

DIMENSIONS = [
    "Technical Skills",
    "Experience & Seniority",
    "Education & Certifications",
    "Domain Knowledge",
    "Soft Skills & Culture Fit",
    "Achievements & Impact",
]

# Max chars to send for resume and JD (~4 chars ≈ 1 token; 6000 chars ≈ 1500 tokens each)
_MAX_RESUME_CHARS = 6_000
_MAX_JD_CHARS    = 5_000

def _truncate(text: str, max_chars: int) -> str:
    """Truncate text to max_chars, appending a notice if truncated."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[...truncated for length...]"

def build_prompt(resume_text: str, jd_text: str, position: str, level: str) -> str:
    resume_text = _truncate(resume_text, _MAX_RESUME_CHARS)
    jd_text     = _truncate(jd_text,     _MAX_JD_CHARS)
    return f"""You are an expert technical recruiter and talent evaluation specialist.

Analyze the following resume against the job description for the position: **{level} {position}**.

Evaluate across these 6 dimensions: {", ".join(DIMENSIONS)}.

For each dimension provide:
- A score from 0 to 10 (float, one decimal place)
- Lists of matched, partially-matched, and missing skills/keywords
- Detailed reasoning (2-3 sentences)

Also provide:
- An overall composite score (0-10)
- Overall reasoning (1-2 sentences)
- Top 3 strengths
- Top 3 gaps
- A final recommendation: exactly one of "Hire", "Maybe", or "Pass"

Respond ONLY with a valid JSON object matching this schema (no markdown fences):
{ANALYSIS_SCHEMA}

---
RESUME:
{resume_text}

---
JOB DESCRIPTION:
{jd_text}
"""

# ── Render dashboard ──────────────────────────────────────────────────────────
def render_dashboard(data: dict, position: str, level: str):
    overall = data.get("overall_score", 0)
    txt_col, bg_col = score_color(overall)

    # ── Recommendation styles (needed for banner) ─────────────────────────────
    rec = data.get("recommendation", "")
    rec_styles = {
        "Hire":  ("background:linear-gradient(135deg,#14532d,#166534);border:1px solid #22c55e;color:#4ade80;", "🚀 Hire"),
        "Maybe": ("background:linear-gradient(135deg,#451a03,#78350f);border:1px solid #f59e0b;color:#fbbf24;", "🤔 Maybe"),
        "Pass":  ("background:linear-gradient(135deg,#450a0a,#7f1d1d);border:1px solid #ef4444;color:#f87171;", "🚫 Pass"),
    }
    rec_style, rec_label = rec_styles.get(rec, ("", rec))

    # ── Overall banner ────────────────────────────────────────────────────────
    st.markdown(f"""
    <div class="overall-banner">
        <div style="font-size:0.85rem;color:#94a3b8;margin-bottom:0.3rem;">
            Fit Score for <strong style="color:#5eead4;">{level} {position}</strong>
        </div>
        <div style="display:flex;align-items:center;justify-content:center;gap:1.2rem;flex-wrap:wrap;margin:0.3rem 0 0.5rem 0;">
            <div class="overall-score-num">{overall:.1f}<span style="font-size:2rem;color:#64748b">/10</span></div>
            <div style="display:flex;flex-direction:row;align-items:center;gap:0.5rem;flex-wrap:wrap;">
                <div class="verdict-chip" style="{verdict_style(overall)}">{verdict_label(overall)}</div>
                <div style="{rec_style} border-radius:99px;padding:5px 16px;font-size:0.85rem;font-weight:700;margin-top: 0.5rem;"
                     title="Recommendation">
                    Recommendation: {rec_label}
                </div>
            </div>
        </div>
        <div class="reasoning-box" style="margin-top:0.8rem;text-align:left;">
            {data.get('overall_reasoning', '')}
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Dimension scores row ──────────────────────────────────────────────────
    dims = data.get("dimensions", [])
    cols = st.columns(len(dims)) if dims else []
    for col, dim in zip(cols, dims):
        sc = dim.get("score", 0)
        tc, bc = score_color(sc)
        with col:
            st.markdown(f"""
            <div class="score-ring-wrap">
                <div class="score-circle" style="color:{tc};border-color:{tc};background:{bc}22;">
                    {sc:.1f}
                </div>
                <div class="score-label">{dim['name']}</div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Dimension detail cards ────────────────────────────────────────────────
    st.markdown('<div class="section-title">📊 Dimension Breakdown</div>', unsafe_allow_html=True)

    for dim in dims:
        sc = dim.get("score", 0)
        tc, _ = score_color(sc)
        grad = meter_gradient(sc)

        matched_html = "".join(f'<span class="badge tag-match">{k}</span>' for k in dim.get("matched", []))
        partial_html = "".join(f'<span class="badge tag-partial">{k}</span>' for k in dim.get("partial", []))
        missing_html = "".join(f'<span class="badge tag-miss">{k}</span>' for k in dim.get("missing", []))

        matched_row = (
            f'<div style="margin-bottom:4px;">'
            f'<span style="font-size:0.75rem;color:#4ade80;font-weight:600;">✅ MATCHED &nbsp;</span>'
            f'{matched_html}</div>'
        ) if matched_html else ""


        partial_row = (
            f'<div style="margin-bottom:4px;">'
            f'<span style="font-size:0.75rem;color:#fbbf24;font-weight:600;">🟡 PARTIAL &nbsp;</span>'
            f'{partial_html}</div>'
        ) if partial_html else ""

        missing_row = (
            f'<div style="margin-bottom:4px;">'
            f'<span style="font-size:0.75rem;color:#f87171;font-weight:600;">❌ MISSING &nbsp;</span>'
            f'{missing_html}</div>'
        ) if missing_html else ""

        tags_block = f'<div style="margin-top:0.8rem;">{matched_row}{partial_row}{missing_row}</div>'
        reasoning = dim.get("reasoning", "")

        st.markdown(f"""
        <div class="glass-card">
            <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:0.5rem;">
                <span style="font-weight:700;font-size:0.95rem;color:#e2e8f0;">{dim['name']}</span>
                <span style="font-weight:800;font-size:1.1rem;color:{tc};">{sc:.1f}/10</span>
            </div>
            <div class="meter-track">
                <div class="meter-fill" style="width:{sc*10}%;background:{grad};"></div>
            </div>
            {tags_block}
            <div class="reasoning-box">{reasoning}</div>
        </div>
        """, unsafe_allow_html=True)


    # ── Strengths & Gaps ──────────────────────────────────────────────────────
    col_s, col_g = st.columns(2)

    with col_s:
        st.markdown('<div class="section-title">💪 Key Strengths</div>', unsafe_allow_html=True)
        # st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        for s in data.get("strengths", []):
            st.markdown(f"✅ &nbsp; {s}")
        st.markdown('</div>', unsafe_allow_html=True)

    with col_g:
        st.markdown('<div class="section-title">⚠️ Gaps to Address</div>', unsafe_allow_html=True)
        # st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        for g in data.get("gaps", []):
            st.markdown(f"🔴 &nbsp; {g}")
        st.markdown('</div>', unsafe_allow_html=True)




# ── Helper: call OpenAI and return parsed analysis dict ──────────────────────
def analyze_resume(client: OpenAI, resume_text: str, jd_text: str,
                   position: str, level: str) -> dict:
    """Run a single resume analysis and return the parsed JSON dict."""
    prompt = build_prompt(resume_text, jd_text, position, level)
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert recruiter. Always respond with valid JSON only."},
            {"role": "user", "content": prompt},
        ],
        temperature=0,
        seed=42,
    )
    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()
    data = json.loads(raw)
    # Enforce score-based recommendation
    s = data.get("overall_score", 0)
    data["recommendation"] = "Hire" if s >= 8 else ("Maybe" if s >= 5 else "Pass")
    return data


# ══════════════════════════════════════════════════════════════════════════════
# MAIN APP
# ══════════════════════════════════════════════════════════════════════════════

# Nav bar
st.markdown("""
<div class="nav-bar">
    <span class="nav-title">📄 Resume Analyzer</span>
    <span style="font-size:0.8rem;color:#64748b;">Powered by GPT-4o-mini</span>
</div>
""", unsafe_allow_html=True)

# Back to Interview button
if st.button("← Back to Interview Helper"):
    st.switch_page("app.py")

st.markdown("### Upload one or more resumes and a job description to get an instant fit analysis.")
st.markdown("---")

# ── Input form ────────────────────────────────────────────────────────────────
with st.form("analyzer_form"):
    col_left, col_right = st.columns([1, 1], gap="large")

    with col_left:
        st.markdown("#### 👤 Resumes")
        resume_files = st.file_uploader(
            "Upload resumes (PDF, DOCX or TXT)",
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True,
            help="Select one or more resumes. In batch mode all files are scored against the same JD.",
        )
        resume_text_input = st.text_area(
            "— or paste a single resume —",
            height=180,
            placeholder="Paste resume text here (single resume only)…",
        )

    with col_right:
        st.markdown("#### 🏢 Job Details")
        level = st.radio(
            "Seniority level",
            options=["Junior", "Mid-level", "Senior", "Lead / Principal"],
            horizontal=True,
        )
        position = st.text_input(
            "Position / Role",
            placeholder="e.g. Machine Learning Engineer",
            max_chars=80,
        )
        jd_text = st.text_area(
            "Job Description",
            height=240,
            placeholder="Paste the full job description here…",
        )

    submitted = st.form_submit_button("🔍 Analyze Fit", use_container_width=True)

# ── Analysis ──────────────────────────────────────────────────────────────────
if submitted:
    # ── Build the list of (filename, resume_text) pairs ──────────────────────
    resume_items: list[tuple[str, str]] = []

    if resume_files:
        for f in resume_files:
            text = extract_text(f)
            if text.strip():
                resume_items.append((f.name, text))
            else:
                st.warning(f"⚠️ Could not extract text from **{f.name}** — skipping.")

    # Fallback to pasted text if no files uploaded
    if not resume_items and resume_text_input.strip():
        resume_items.append(("Pasted Resume", resume_text_input.strip()))

    # Validate
    errors = []
    if not resume_items:
        errors.append("Please upload at least one resume (or paste resume text).")
    if not jd_text.strip():
        errors.append("Please provide a job description.")
    if not position.strip():
        errors.append("Please specify the position/role.")

    if errors:
        for e in errors:
            st.error(e)
        st.stop()

    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
    is_batch = len(resume_items) > 1

    # ── Run analysis for each resume ─────────────────────────────────────────
    results: list[tuple[str, dict]] = []   # (filename, data)
    progress = st.progress(0, text="Preparing analysis…")

    for idx, (fname, rtext) in enumerate(resume_items):
        progress.progress(
            (idx) / len(resume_items),
            text=f"🤖 Analyzing **{fname}** ({idx + 1}/{len(resume_items)})…"
        )
        try:
            data = analyze_resume(client, rtext, jd_text.strip(), position.strip(), level)
            results.append((fname, data))
        except json.JSONDecodeError as e:
            st.error(f"**{fname}**: Could not parse model response — {e}")
        except Exception as e:
            st.error(f"**{fname}**: API error — {e}")

    progress.progress(1.0, text="✅ Analysis complete!")

    if not results:
        st.stop()

    st.markdown("---")

    # ── Batch: ranked leaderboard table ──────────────────────────────────────
    if is_batch:
        # Sort by overall score descending
        results.sort(key=lambda x: x[1].get("overall_score", 0), reverse=True)

        MEDALS = ["🥇", "🥈", "🥉"]
        REC_COLORS = {
            "Hire":  ("#4ade80", "rgba(34,197,94,0.15)"),
            "Maybe": ("#fbbf24", "rgba(251,191,36,0.15)"),
            "Pass":  ("#f87171", "rgba(239,68,68,0.15)"),
        }

        st.markdown(
            f'<div class="section-title">📋 Candidate Ranking — {len(results)} Resumes'
            f' · {position.strip()} ({level})</div>',
            unsafe_allow_html=True
        )

        rows_html = ""
        for rank, (fname, d) in enumerate(results):
            score = d.get("overall_score", 0)
            rec   = d.get("recommendation", "")
            txt_c, bg_c = REC_COLORS.get(rec, ("#e2e8f0", "rgba(255,255,255,0.05)"))
            medal = MEDALS[rank] if rank < 3 else f"#{rank + 1}"
            tc, _ = score_color(score)

            rows_html += f"""
            <tr>
              <td><span class="rank-medal">{medal}</span></td>
              <td style="font-weight:600;color:#e2e8f0;">{fname}</td>
              <td style="font-weight:800;color:{tc};font-size:1.05rem;">{score:.1f}<span style="font-size:0.8rem;color:#64748b;">/10</span></td>
              <td><span style="background:{bg_c};color:{txt_c};border:1px solid {txt_c}44;
                           border-radius:99px;padding:3px 12px;font-size:0.78rem;font-weight:700;">
                {rec}</span></td>
              <td style="color:#94a3b8;font-size:0.82rem;max-width:320px;">
                {', '.join(d.get('strengths', [])[:2])}
              </td>
            </tr>"""

        st.markdown(f"""
        <div class="glass-card" style="padding:1rem 1.2rem;">
          <table class="batch-table">
            <thead>
              <tr>
                <th>Rank</th><th>Candidate File</th><th>Score</th>
                <th>Decision</th><th>Top Strengths</th>
              </tr>
            </thead>
            <tbody>{rows_html}</tbody>
          </table>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-title">🔍 Individual Dashboards</div>', unsafe_allow_html=True)

        for rank, (fname, d) in enumerate(results):
            score = d.get("overall_score", 0)
            rec   = d.get("recommendation", "")
            medal = MEDALS[rank] if rank < 3 else f"#{rank + 1}"
            tc, _ = score_color(score)
            with st.expander(
                f"{medal}  {fname}  ·  {score:.1f}/10  ·  {rec}",
                expanded=(rank == 0),   # auto-expand top candidate
            ):
                render_dashboard(d, position.strip(), level)

    # ── Single resume: original full-page dashboard ───────────────────────────
    else:
        fname, data = results[0]
        render_dashboard(data, position.strip(), level)
